from crewai import Agent, Task, Crew, Process
from typing import Dict, List
from agents.senior_agent import SeniorAgent
from utils.file_utils import extract_zip, cleanup_directory
from utils.analysis_utils import analyze_folder
from docx import Document
import asyncio

class SuperSeniorAgent:
    def __init__(self):
        self.super_agent = Agent(
            role='Super Senior Developer and Technical Architect',
            goal=(
                """Oversee comprehensive technical analysis and documentation with these objectives:
                1. Examine code architecture, patterns, and design decisions.
                2. Assess code quality, maintainability, and scalability.
                3. Identify technical debt and propose improvements.
                4. Produce detailed technical documentation:
                   - Architecture diagrams (Mermaid)
                   - Data flow and sequence diagrams
                   - Component interaction maps
                   - Code quality assessments
                   - Improvement recommendations
                5. Synthesize findings from all Senior Agents into actionable insights."""
            ),
            backstory=(
                """A seasoned technical architect with 15+ years of enterprise software experience, skilled in:
                - Analyzing complex distributed systems
                - Identifying architectural patterns and anti-patterns
                - Creating robust technical documentation
                - Defining improvement strategies
                - Mentoring development teams
                - Translating technical complexity into clear deliverables

                Known for transforming intricate codebases into maintainable, well-documented solutions."""
            ),
            llm="openrouter/google/gemini-flash-1.5",
        )
    async def execute(self, zip_file_path: str):
        try:
            extracted_folder = extract_zip(zip_file_path)
            language_files = self.analyze_folder(extracted_folder)
            if not language_files:
                print("No code files found in the zip file.")
                return "No code files found."

            senior_agents = self.create_senior_agents(language_files)
            senior_summaries = await asyncio.gather(*(senior_agent.execute() for senior_agent in senior_agents))
            # Create and return the final summary
            return self.create_final_summary(senior_summaries)
        finally:
            cleanup_directory(extracted_folder)
    
    def analyze_folder(self, folder_path: str) -> Dict[str, List[str]]:
        return analyze_folder(folder_path)
    

    def create_senior_agents(self, language_files: Dict[str, List[str]]) -> List['SeniorAgent']:
        senior_agents = []
        for language, files in language_files.items():
            senior_agent = SeniorAgent(language, files)
            senior_agents.append(senior_agent)
        return senior_agents    

    def create_final_summary(self, senior_summaries: List[str]):
        description = (
            "As the Super Senior Developer, compile a detailed technical document for the entire project based on the following Senior Agents' summaries:\n\n"
            + "\n\n".join(senior_summaries)
            + "\n\nEnsure the documentation covers key architecture, data flows, notable findings, and includes Mermaid diagram code for visualization."
        )
        expected_output = (
            "A comprehensive technical document detailing the project's architecture, data flows, and code assessments, "
            "including Mermaid diagrams for visualization."
        )

        task = Task(
            description=description,
            expected_output=expected_output,
            agent=self.super_agent,
        )
        crew = Crew(
            agents=[self.super_agent],
            tasks=[task],
            process=Process.sequential
        )
        output = crew.kickoff()        # Save to a .docx file
        self.save_to_docx(output)

        # Return the content for Streamlit display
        return output.content if hasattr(output, 'content') else str(output)

    def save_to_docx(self, output):
        doc = Document()
        doc.add_heading('Final Technical Documentation', level=1)
        if hasattr(output, 'content'):
            doc.add_paragraph(output.content)
        else:
            doc.add_paragraph(str(output))

        try:
            doc.save('technical_documentation.docx')
            print("Technical documentation created successfully.")
        except Exception as e:
            print(f"An error occurred while saving the document: {e}")
